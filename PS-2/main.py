import pandas as pd
import docx
import argparse
import sys
import os
from experiment import Experiment
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
# from flask import Flask, render_template, send_file, flash, request, redirect, url_for, Response
# from werkzeug.utils import secure_filename
from tkinter import filedialog
from tkinter import *
from tkinter import ttk

global in_file, in_temp_file, out_file 
def getFilePath(type, file):
    global in_file, in_temp_file
    val = filedialog.askopenfilename(
        initialdir="/", title="Select file", filetypes=type)
    if file == 'doc':
        in_temp_file = val
    else:
        in_file = val


def getSavePath():
    global out_file, in_file, in_temp_file
    out_file = filedialog.asksaveasfile(mode='w', defaultextension=".docx")
    print(in_temp_file, out_file)
    # import pdb; pdb.set_trace()
    out_file = str(out_file.name)
    convert(None, in_file, in_temp_file, out_file)


FONT = ("Myriad Pro", 13)
LARGE_FONT = ("Myriad Pro", 30)

keys = {'Exp. Title': None, 'Exp. reference ': [1, 1], 'Author Comments': [1, 2], 'Formula Used': [1, 4], 'TestGoal': [1, 9], 'TestMeans': None,
        'Date': [4, 1], 'Author': [4, 2], 'TestMethod': None, 'Testtype': None, 'ExpStatus': None, 'Conclusion': None}

exp_method_options = {'Select Test Method',
                      'Aotomatic', 'Semi-Automatic', 'Manual'}
exp_status_options = {'Select Test Status', 'InProgress', 'Done', 'Reviewed'}
exp_type_options = {'Select Test Type', 'Type1', 'Type2', 'Type3'}
exp_conclusion_options = {'OK', 'NOK', 'NOT DONE'}


def convert(args=None, file_path=None, template_path=None, out_path=None):
    if args is not None:
        if args.file == None and args.template == None:
            print('ERROR :: Plesae specify path for excel file and template file\n\n')
            argparse.ArgumentParser.print_help(arg_parser)
            sys.exit(1)

        if '.docx' != args.template[-5:] or '.xlsx' != args.file[-5:] or ('.docx' != args.output[-5:] if args.output != None else False):
            print('ERROR :: Invalid input file type\n\n')
            argparse.ArgumentParser.print_help(arg_parser)
            sys.exit(1)

        file_path = os.path.expanduser(args.file)
        template_path = args.template
        out_path = args.output if args.output != None else file_path.split('.')[
            0] + '.docx'

    df = pd.read_excel(file_path)
    template_doc = docx.Document(template_path)

    experiments = list(df['Experiment ID'].unique())
    parsed = {}
    for exp in experiments:
        if str(exp) != 'nan':
            experiment = Experiment(exp, df, keys.keys())
            parsed[exp] = experiment.parsed_data

    titles = {}
    for para in template_doc.paragraphs:
        if 'title' in para.text:
            titles[para] = None

    if len(template_doc.tables) != len(titles.keys())*2:
        print('ERROR :: The number of tables in the template does not corrosponds to number of experiment titles\nNumber of Titles :{}\nNumber of Expected Tables :{}\nNumber of Actual Tables :{}'.format(
            len(titles), len(titles)*2, len(template_doc.tables)))
    for title, i in zip(titles.keys(), range(0, len(titles))):
        titles[title] = list(
            (template_doc.tables[2 * i], template_doc.tables[2 * i + 1]))
    s_no = 1
    for title in titles:
        exp_num = int(title.text.split()[-1])
        exp = 'EXP_' + '0' * (3 - len(str(exp_num))) + str(exp_num)
        for key in keys.keys():
            # print('now', key)
            if keys[key] is None:
                continue
            try:
                val = parsed[exp][key] if not isinstance(
                    parsed[exp][key], pd._libs.tslibs.timestamps.Timestamp) else (str(parsed[exp][key].day) + '/' + str(parsed[exp][key].month) + '/' + str(parsed[exp][key].year))
                # print(val)
                titles[title][0].column_cells(keys[key][0])[
                    keys[key][1]].text = val
            except:
                print('Error Occured', title, key,
                      keys[key][0], keys[key][1], parsed[exp][key])
        description_summary = ''
        for p in parsed[exp]['Procedure']:
            description_summary += '-' + str(p[0]) + '\n'
        # print(description_summary)

        titles[title][0].column_cells(1)[10].text = ''
        desc_para = titles[title][0].column_cells(
            1)[10].add_paragraph(description_summary)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # print(titles[title][1].row_cells(0)[0].text)
        titles[title][1].autofit = True
        titles[title][1].alignment = WD_TABLE_ALIGNMENT.LEFT
        j = 0
        total_rows = len(titles[title][1].rows)
        for p, i in zip(parsed[exp]['Procedure'], range(len(parsed[exp]['Procedure']))):
            if j < total_rows-2:
                # titles[title][1].row_cells(i + 1)[0].text = str(s_no)
                # print(p[0])
                # print(i)
                titles[title][1].row_cells(i + 1)[2].text = p[0]
                titles[title][1].row_cells(
                    i + 1)[2].alignment = WD_ALIGN_PARAGRAPH.LEFT
                titles[title][1].row_cells(
                    i + 1)[3].text = p[1] if p[1] is not None else ''
                titles[title][1].row_cells(
                    i + 1)[3].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                row = titles[title][1].add_row()
                # import pdb; pdb.set_trace()
                row.cells[0].text = str(s_no)
                row.cells[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                # row.cells[0].merge(row.cells[1])
                row.cells[1].text = p[0]
                row.cells[2].text = p[1] if p[1] is not None else ''
                row.cells[3].text = ''
                row.cells[4].text = ''
                row.cells[3].merge(row.cells[4])
                titles[title][1].rows[j]._tr.addnext(row._tr)
                # print('IndexError')
            j += 1
            s_no += 1
        title.text = parsed[exp]['Exp. Title']
        titles[title][0].column_cells(2)[6].text = ''
        titles[title][0].column_cells(2)[7].text = ''
        titles[title][0].column_cells(2)[8].text = ''
        titles[title][0].column_cells(4)[7].text = ''
    template_doc.save(out_path)


ALLOWED_EXTENSIONS = {'docx', 'xlsx'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
        
if __name__ == '__main__':
    arg_parser = argparse.ArgumentParser(
        description='Convert Excel Sheet(.xlsx) to Word Document(.docx) based on Template')

    arg_parser.add_argument('-f', '--file',
                            help='Excel File Path (.xlsx)',
                            required=False,
                            default=None)

    arg_parser.add_argument('-t', '--template',
                            help='Template for Conversion (.docx)',
                            required=False,
                            default=None)

    arg_parser.add_argument('-o', '--output',
                            help='Output File Name (.docx)',
                            required=False,
                            default=None)
    arg_parser.add_argument('-g', '--gui',
                            help='Start in GUI Mode',
                            required=False,
                            default=None)

    args = sys.argv[1:]
    args = arg_parser.parse_args(args)
    if args.gui == None:
        convert(args)
    else:
        root = Tk()
        l1 = Label(root, text="Excel to Docx Parser", font=LARGE_FONT)
        l1.grid(row=0, column=1, sticky=W, pady=20)

        root.geometry('600x300')
        template_btn = ttk.Button(
            root, text="Choose Template", command=lambda: getFilePath((("docx file", "*.docx"), ("all files", "*.*")), 'doc'), width=10)
        excel_btn = ttk.Button(
            root, text="Choose Excel File", command=lambda: getFilePath((("excel sheet", "*.xlsx"), ("all files", "*.*")), 'xl'), width=10)
        parse_btn = ttk.Button(
            root, text="Parse", command=lambda: getSavePath(), width=10)
        template_btn.grid(row=1, column=0, sticky=W, pady=100, padx=10)
        excel_btn.grid(row=1, column=1, sticky=W, pady=100, padx=100)
        parse_btn.grid(row=1, column=2, sticky=W, pady=100)
        # template_btn.pack(side='top')
        root.mainloop()
    
